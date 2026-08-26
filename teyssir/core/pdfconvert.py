"""PDF → Word (.docx) conversion — free/OSS, Windows-hub optimized.

Two engines, auto-selected by content analysis:

* **fast**   — text-dense PDFs: PyMuPDF text extract → python-docx (10–50× faster).
* **layout** — mixed/graphic PDFs: pdf2docx with tuned kwargs (clip 2.0, no stream
  tables, optional multi-processing for ≥8 pages).

Temp files live under ``MEDIA_ROOT/tmp`` (not system ``%TEMP%``) so Windows Defender
and redirected profiles do not scan every write. Prefer ``Converter(stream=…)`` when
not multiprocessing to avoid a temp PDF write.
"""
from __future__ import annotations

import logging
import os
import tempfile
from dataclasses import dataclass
from io import BytesIO

MAX_PDF_BYTES = 25 * 1024 * 1024        # 25 MB guard
# Sync HTTP path reserved for tiny jobs (keeps backward-compatible 200 + FileResponse).
SYNC_MAX_BYTES = 2 * 1024 * 1024
SYNC_MAX_PAGES = 5
# Text-density heuristics for the fast path.
FAST_MIN_CHARS_PER_PAGE = 80
FAST_MAX_IMAGES_PER_PAGE = 1.5
FAST_MAX_PAGES = 80                     # beyond this, prefer layout engine for fidelity

# pdf2docx is extremely chatty at INFO; keep service logs readable.
logging.getLogger("pdf2docx").setLevel(logging.WARNING)
log = logging.getLogger("teyssir.pdfconvert")


@dataclass
class PdfProfile:
    pages: int
    chars: int
    images: int
    size_bytes: int

    @property
    def avg_chars(self) -> float:
        return self.chars / self.pages if self.pages else 0.0

    @property
    def avg_images(self) -> float:
        return self.images / self.pages if self.pages else 0.0

    @property
    def is_text_dense(self) -> bool:
        return (
            self.pages > 0
            and self.pages <= FAST_MAX_PAGES
            and self.avg_chars >= FAST_MIN_CHARS_PER_PAGE
            and self.avg_images <= FAST_MAX_IMAGES_PER_PAGE
        )

    @property
    def fits_sync(self) -> bool:
        return self.size_bytes <= SYNC_MAX_BYTES and self.pages <= SYNC_MAX_PAGES


def media_tmp_dir() -> str:
    """Prefer MEDIA_ROOT/tmp on the shop data volume (Windows AV-friendly) over %TEMP%."""
    try:
        from django.conf import settings
        root = getattr(settings, "MEDIA_ROOT", None)
    except Exception:  # noqa: BLE001 — settings may be unavailable in unit probes
        root = None
    if root:
        path = os.path.join(str(root), "tmp")
        os.makedirs(path, exist_ok=True)
        return path
    return tempfile.gettempdir()


def convert_workspace(job_id) -> str:
    """Per-job directory under MEDIA_ROOT/convert/<job_id>/."""
    from django.conf import settings

    path = os.path.join(str(settings.MEDIA_ROOT), "convert", str(job_id))
    os.makedirs(path, exist_ok=True)
    return path


def validate_pdf_header(pdf_bytes: bytes) -> None:
    if not pdf_bytes:
        raise ValueError("empty file")
    if len(pdf_bytes) > MAX_PDF_BYTES:
        raise ValueError("PDF larger than 25 MB")
    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("not a PDF file")


def profile_pdf(pdf_bytes: bytes) -> PdfProfile:
    """Cheap PyMuPDF pass: page count, text density, image count."""
    import fitz

    validate_pdf_header(pdf_bytes)
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        chars = images = 0
        for page in doc:
            chars += len(page.get_text("text") or "")
            images += len(page.get_images(full=True) or [])
        return PdfProfile(
            pages=doc.page_count, chars=chars, images=images, size_bytes=len(pdf_bytes),
        )
    finally:
        doc.close()


def choose_mode(profile: PdfProfile, requested: str = "auto") -> str:
    if requested in ("fast", "layout"):
        return requested
    return "fast" if profile.is_text_dense else "layout"


def _convert_fast(pdf_bytes: bytes) -> bytes:
    """Text path: extract page text into a simple .docx (editable, layout-light)."""
    import fitz
    from docx import Document

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        wd = Document()
        for i, page in enumerate(doc):
            if i:
                wd.add_page_break()
            text = (page.get_text("text") or "").strip()
            if text:
                for para in text.split("\n"):
                    wd.add_paragraph(para)
            else:
                wd.add_paragraph(f"[Page {i + 1} — no extractable text]")
        buf = BytesIO()
        wd.save(buf)
        return buf.getvalue()
    finally:
        doc.close()


def _convert_layout(pdf_bytes: bytes, *, pages: int) -> bytes:
    """Layout-preserving path via pdf2docx with Windows-tuned defaults."""
    from pdf2docx import Converter

    use_mp = pages >= 8
    cpu = min(4, os.cpu_count() or 1) if use_mp else 0
    kwargs = dict(
        clip_image_res_ratio=2.0,
        parse_stream_table=False,
        parse_lattice_table=True,
        ignore_page_error=True,
    )
    if use_mp:
        kwargs.update(multi_processing=True, cpu_count=cpu, start=0, end=pages)

    # Multiprocessing re-opens a real path; stream= avoids the temp PDF write otherwise.
    with tempfile.TemporaryDirectory(dir=media_tmp_dir()) as tmp:
        dst = os.path.join(tmp, "out.docx")
        if use_mp:
            src = os.path.join(tmp, "in.pdf")
            with open(src, "wb") as fh:
                fh.write(pdf_bytes)
            converter = Converter(src)
        else:
            converter = Converter(stream=pdf_bytes)
        try:
            converter.convert(dst, **kwargs)
        finally:
            converter.close()
        with open(dst, "rb") as fh:
            return fh.read()


def convert_pdf_to_docx(pdf_bytes, *, mode: str = "auto") -> tuple[bytes, str, PdfProfile]:
    """Convert PDF bytes → DOCX bytes.

    Returns ``(docx_bytes, mode_used, profile)``. Raises ``ValueError`` on bad input.
    """
    profile = profile_pdf(pdf_bytes)
    used = choose_mode(profile, mode)
    log.info(
        "pdfconvert pages=%s chars=%s images=%s mode=%s size=%s",
        profile.pages, profile.chars, profile.images, used, profile.size_bytes,
    )
    if used == "fast":
        return _convert_fast(pdf_bytes), used, profile
    return _convert_layout(pdf_bytes, pages=profile.pages), used, profile


def convert_pdf_file(src_path: str, dst_path: str, *, mode: str = "auto") -> tuple[str, PdfProfile]:
    """Convert a PDF on disk to a DOCX on disk. Returns ``(mode_used, profile)``."""
    with open(src_path, "rb") as fh:
        pdf_bytes = fh.read()
    docx, used, profile = convert_pdf_to_docx(pdf_bytes, mode=mode)
    os.makedirs(os.path.dirname(dst_path) or ".", exist_ok=True)
    with open(dst_path, "wb") as fh:
        fh.write(docx)
    return used, profile
