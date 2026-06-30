#!/usr/bin/env python
"""Convert docs/ARCHITECTURE.md -> a stakeholder DOCX (python-docx).

Handles headings, paragraphs (with **bold** / *italic* / `code` / [links]), pipe
tables, fenced code blocks (incl. mermaid/ASCII as monospace), bullet/numbered lists,
blockquotes and horizontal rules. Soft-wrapped Markdown lines are joined into one
logical block so inline spans that wrap across lines render correctly. Diagrams are
kept as monospace source — view the original Markdown for the rendered visuals. Usage:

    python tools/md_to_docx.py docs/ARCHITECTURE.md docs/ARCHITECTURE.docx
"""
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+?\*|\[[^\]]+\]\([^)]+\))")
ACCENT = RGBColor(0x1B, 0x5E, 0x20)   # library green (spec §11.3)
GREY = RGBColor(0x55, 0x55, 0x55)
CODE_FILL = "F2F2F2"
HEAD_FILL = "1B5E20"


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _tokens(text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            yield text[pos:m.start()], {}
        tok = m.group(0)
        if tok.startswith("**"):
            yield tok[2:-2], {"bold": True}
        elif tok.startswith("`"):
            yield tok[1:-1], {"code": True}
        elif tok.startswith("*"):
            yield tok[1:-1], {"italic": True}
        else:
            yield re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok).group(1), {}
        pos = m.end()
    if pos < len(text):
        yield text[pos:], {}


def add_inline(paragraph, text, *, base_bold=False, base_italic=False, color=None, size=None):
    for t, fmt in _tokens(text):
        run = paragraph.add_run(t)
        run.bold = bool(fmt.get("bold") or base_bold)
        run.italic = bool(fmt.get("italic") or base_italic)
        if fmt.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        if size is not None:
            run.font.size = size
        if color is not None:
            run.font.color.rgb = color


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, CODE_FILL)
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_table(doc, rows):
    header, body = rows[0], rows[2:]  # rows[1] is the |---| separator
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True
    for j, cellt in enumerate(header):
        c = table.cell(0, j)
        shade(c, HEAD_FILL)
        add_inline(c.paragraphs[0], cellt, base_bold=True,
                   color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))
    for row in body:
        cells = table.add_row().cells
        for j in range(len(header)):
            add_inline(cells[j].paragraphs[0], row[j] if j < len(row) else "", size=Pt(9))


def convert(md_path, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(20)
    sec.left_margin = sec.right_margin = Mm(18)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for lvl in range(1, 5):
        doc.styles[f"Heading {lvl}"].font.color.rgb = ACCENT

    pending = {"kind": None, "text": "", "indent": False}

    def flush():
        kind, text, indent = pending["kind"], pending["text"], pending["indent"]
        pending.update(kind=None, text="", indent=False)
        if kind is None:
            return
        if kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            if indent:
                p.paragraph_format.left_indent = Mm(14)
            add_inline(p, text)
        elif kind == "number":  # keep source numbering (avoid python-docx auto-renumber)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(8)
            p.paragraph_format.first_line_indent = Mm(-5)
            add_inline(p, text)
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(6)
            add_inline(p, text, base_italic=True, color=GREY)
        else:  # para
            add_inline(doc.add_paragraph(), text)

    def hr():
        p = doc.add_paragraph()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:color", "1B5E20"), ("w:space", "1")):
            bottom.set(qn(k), v)
        pbdr.append(bottom)
        p._p.get_or_add_pPr().append(pbdr)

    lines = open(md_path, encoding="utf-8").read().split("\n")
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        if line.lstrip().startswith("```"):
            flush()
            lang = line.strip().strip("`").strip()
            block, i = [], i + 1
            while i < n and not lines[i].lstrip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            if lang == "mermaid":
                add_inline(doc.add_paragraph(),
                           "*Diagram (Mermaid source — see ARCHITECTURE.md for the rendered visual):*")
            add_code_block(doc, block)
            continue

        if line.strip().startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|:- "):
            flush()
            rows, j = [], i
            while j < n and lines[j].strip().startswith("|"):
                rows.append(split_row(lines[j]))
                j += 1
            add_table(doc, rows)
            i = j
            continue

        stripped = line.strip()

        if stripped.startswith("#"):
            flush()
            m = re.match(r"(#+)\s+(.*)", stripped)
            level = len(m.group(1))
            text = re.sub(r"`([^`]+)`", r"\1", m.group(2))
            if level == 1:
                doc.add_heading(text, 0)
            else:
                add_inline(doc.add_heading("", min(level - 1, 4)), text)
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            flush()
            hr()
            i += 1
            continue

        if not stripped:
            flush()
            i += 1
            continue

        mbul = re.match(r"^(\s*)[-*]\s+(.*)", line)
        mnum = re.match(r"^(\s*)(\d+\.)\s+(.*)", line)
        if mbul:
            flush()
            pending.update(kind="bullet", text=mbul.group(2), indent=len(mbul.group(1)) >= 2)
        elif mnum:
            flush()
            pending.update(kind="number", text=f"{mnum.group(2)} {mnum.group(3)}",
                           indent=len(mnum.group(1)) >= 2)
        elif stripped.startswith(">"):
            qt = stripped.lstrip("> ").rstrip()
            if pending["kind"] == "quote":
                pending["text"] += " " + qt
            else:
                flush()
                pending.update(kind="quote", text=qt)
        elif pending["kind"] in ("para", "bullet", "number", "quote"):
            pending["text"] += " " + stripped       # join soft-wrapped continuation
        else:
            pending.update(kind="para", text=stripped)
        i += 1

    flush()
    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
